from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DOCX_PATH = ROOT / "AlarmMini_User_Manual_uk.docx"
TELEGRAM_URL = "https://t.me/+j3zFZHE5gGoyNGYy"

BLACK = RGBColor(0x17, 0x1A, 0x1C)
GRAY = RGBColor(0x58, 0x63, 0x6D)
LIGHT_GRAY = RGBColor(0xD7, 0xDF, 0xE6)
SOFT_BLUE = RGBColor(0xE8, 0xEE, 0xF5)
BLUE = RGBColor(0x1F, 0x4D, 0x78)
GOLD = RGBColor(0x9A, 0x6A, 0x00)
INK = RGBColor(0x2C, 0x3E, 0x50)
MIST = RGBColor(0xF6, 0xF8, 0xFB)
PALE_GOLD = RGBColor(0xF7, 0xF1, 0xE3)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D7DFE6", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        el = tc_borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def set_run_font(run, size, color=BLACK, bold=False, italic=False, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(doc, text, *, style=None, size=11, color=BLACK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.25):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=line)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, before=0, after=4, line=1.25)
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_numbers(doc, items, style_name="List Number"):
    for item in items:
        p = doc.add_paragraph(style=style_name)
        set_paragraph_spacing(p, before=0, after=4, line=1.25)
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.rows[0].cells[0]
    set_cell_width(cell, 6.5)
    set_cell_shading(cell, "F4F6F9")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = cell.paragraphs[0]
    set_paragraph_spacing(p1, before=0, after=2, line=1.15)
    r1 = p1.add_run(title)
    set_run_font(r1, size=10.5, color=BLUE, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=0, line=1.2)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5, color=BLACK)
    doc.add_paragraph()


def add_banner(doc, title, body, *, fill="1F4D78", title_color=SOFT_BLUE, body_color=MIST):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.rows[0].cells[0]
    set_cell_width(cell, 6.5)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=fill, size="0")
    p1 = cell.paragraphs[0]
    set_paragraph_spacing(p1, before=0, after=4, line=1.0)
    r1 = p1.add_run(title)
    set_run_font(r1, size=12, color=title_color, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=0, line=1.15)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=body_color)
    doc.add_paragraph()


def add_meta_strip(doc, items):
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [6.5 / len(items)] * len(items)
    for idx, (label, value) in enumerate(items):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "F6F8FB")
        set_cell_border(cell, color="D7DFE6", size="6")
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p1, before=0, after=2, line=1.0)
        r1 = p1.add_run(label)
        set_run_font(r1, size=9.5, color=BLUE, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, before=0, after=0, line=1.1)
        r2 = p2.add_run(value)
        set_run_font(r2, size=10.5, color=INK, bold=True)
    doc.add_paragraph()


def add_section_lead(doc, eyebrow, title, body):
    add_text(
        doc,
        eyebrow,
        size=9.5,
        color=GOLD,
        bold=True,
        before=10,
        after=2,
        line=1.0,
    )
    add_text(
        doc,
        title,
        size=21,
        color=BLUE,
        bold=True,
        before=0,
        after=5,
        line=1.0,
    )
    add_text(
        doc,
        body,
        size=11.5,
        color=INK,
        before=0,
        after=10,
        line=1.25,
    )


def add_two_col_table(doc, rows, widths=(1.8, 4.7), header=None):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if header:
        hdr = table.add_row().cells
        for idx, text in enumerate(header):
            set_cell_width(hdr[idx], widths[idx])
            set_cell_shading(hdr[idx], "E8EEF5")
            p = hdr[idx].paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=10.5, bold=True, color=BLUE)
    for left, right in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((left, right)):
            set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.2)
            r = p.add_run(text)
            set_run_font(r, size=10.5, bold=(idx == 0 and header is None))
    doc.add_paragraph()


def add_three_col_table(doc, rows, widths=(1.5, 2.2, 2.8), header=None):
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if header:
        hdr = table.add_row().cells
        for idx, text in enumerate(header):
            set_cell_width(hdr[idx], widths[idx])
            set_cell_shading(hdr[idx], "E8EEF5")
            p = hdr[idx].paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=10.5, bold=True, color=BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            p = cells[idx].paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.2)
            r = p.add_run(text)
            set_run_font(r, size=10.25)
    doc.add_paragraph()


def add_heading(doc, text, level):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        set_paragraph_spacing(p, before=18, after=8, line=1.15)
        size, color = 16, BLUE
    elif level == 2:
        set_paragraph_spacing(p, before=14, after=7, line=1.15)
        size, color = 13, BLUE
    else:
        set_paragraph_spacing(p, before=10, after=5, line=1.1)
        size, color = 12, RGBColor(0x1F, 0x3A, 0x5F)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)
    return p


def set_base_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)


def add_footer(section):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    for para in list(footer.paragraphs):
        para.clear()


def add_cover(doc):
    add_banner(
        doc,
        "ALARMMINI USER GUIDE",
        "Покрокова інструкція для першого запуску, підключення до Wi-Fi, прошивки та базового налаштування карти.",
    )
    add_text(
        doc,
        "AlarmMini",
        size=30,
        color=BLUE,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=28,
        after=8,
        line=1.0,
    )
    add_text(
        doc,
        "Підключення, прошивка та базове налаштування карти тривог України",
        size=14.5,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=16,
        line=1.2,
    )
    add_meta_strip(
        doc,
        [
            ("Формат", "Коротко і по кроках"),
            ("Перший запуск", "Через Wi-Fi AP"),
            ("Оновлення", "Через USB і сайт"),
        ],
    )
    add_two_col_table(
        doc,
        [
            ("Сумісні плати", "ESP32-C3 SuperMini, ESP8266 / Wemos D1 mini"),
            ("Перше що шукати", "Wi-Fi мережу виду AlarmMap-Setup-XXXXXX"),
            ("Сторінка налаштування", "http://192.168.4.1"),
            ("Пароль точки доступу", "Без пароля"),
            ("Сайт для прошивки", "https://alarmmini.vercel.app"),
        ],
        widths=(1.9, 4.6),
    )
    add_note(
        doc,
        "Почніть саме з цього",
        "Після ввімкнення спочатку знайдіть у списку Wi-Fi мереж точку доступу виду AlarmMap-Setup-XXXXXX і підключіться до неї. Лише після цього відкривайте 192.168.4.1 або переходьте до прошивки та інших налаштувань.",
    )
    add_text(
        doc,
        f"Якщо потрібна допомога, автору можна написати в Telegram: {TELEGRAM_URL}",
        size=10,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=4,
        after=0,
        line=1.15,
    )


def build_doc():
    doc = Document()
    set_base_styles(doc)
    add_footer(doc.sections[0])
    add_cover(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer(doc.sections[-1])

    add_section_lead(
        doc,
        "ПЕРШИЙ КРОК",
        "1. Перше підключення до Wi-Fi",
        "Якщо пристрій ще не має збережених Wi-Fi даних, він автоматично створює власну точку доступу. Саме з цього сценарію користувач має починати роботу.",
    )
    add_two_col_table(
        doc,
        [
            ("SSID", "AlarmMap-Setup-XXXXXX"),
            ("Що означає XXXXXX", "Унікальний суфікс вашої плати, тому останні символи у різних пристроїв відрізняються"),
            ("Пароль AP", "Без пароля"),
            ("Адреса сторінки налаштування", "http://192.168.4.1"),
            ("Коли відкривати сайт прошивки", "Після підключення до цієї Wi-Fi мережі або коли потрібно оновити прошивку через USB"),
        ],
        widths=(2.1, 4.4),
    )
    add_bullets(
        doc,
        [
            "У списку Wi-Fi мереж знайдіть мережу виду AlarmMap-Setup-XXXXXX. Закінчення XXXXXX у вашого пристрою буде своїм.",
            "Підключіться до цієї мережі. Пароль не потрібен.",
            "Якщо сторінка не відкрилась автоматично, вручну відкрийте 192.168.4.1 у браузері.",
            "Оберіть вашу Wi-Fi мережу 2.4 ГГц і введіть пароль.",
            "Натисніть «Save» або підтвердіть збереження налаштувань.",
            "Після перезавантаження поверніться у свою звичайну Wi-Fi мережу та продовжуйте налаштування.",
        ],
    )
    add_note(
        doc,
        "Не шукайте сайт одразу",
        "Якщо ви просто вперше вмикаєте пристрій, не потрібно відразу переходити на сайт прошивки. Спочатку підключіться до Wi-Fi мережі AlarmMap-Setup-XXXXXX і завершіть базове налаштування мережі.",
    )

    add_heading(doc, "2. Що знадобиться", 1)
    add_bullets(
        doc,
        [
            "Пристрій AlarmMini, підключений до живлення або до комп'ютера через USB.",
            "Якісний USB-кабель для передачі даних. Кабель лише для зарядки не підійде.",
            "Комп'ютер або ноутбук з браузером Chrome чи Microsoft Edge. Саме ці браузери потрібні для Web Serial та прошивки.",
            "Домашня або офісна Wi-Fi мережа 2.4 ГГц. Мережі 5 ГГц пристроєм не підтримуються.",
            "Дані вашого MQTT-брокера: адреса, порт, topic, за потреби логін і пароль.",
        ],
    )
    add_note(
        doc,
        "Важливо",
        "Якщо ви плануєте лише перше підключення до Wi-Fi через режим точки доступу, смартфона достатньо. Для прошивки через сайт краще використовувати комп'ютер.",
    )

    add_banner(
        doc,
        "Швидкий маршрут",
        "Якщо хочете пройти весь шлях без зайвих деталей: спочатку Wi-Fi точка доступу, потім за потреби прошивка через сайт, далі MQTT, web-панель і калібрування.",
        fill="F7F1E3",
        title_color=GOLD,
        body_color=INK,
    )
    add_heading(doc, "3. Швидкий старт", 1)
    add_numbers(
        doc,
        [
            "Якщо це перший запуск, спочатку підключіться до Wi-Fi мережі AlarmMap-Setup-XXXXXX і збережіть вашу домашню мережу.",
            "Для прошивки підключіть плату до комп'ютера через USB.",
            "Відкрийте сайт https://alarmmini.vercel.app у Chrome або Edge.",
            "Натисніть «Підключити плату» та дозвольте браузеру доступ до серійного порту.",
            "Оберіть тип плати та прошийте пристрій: «Прошити як новий пристрій» або «Прошити з відновленням налаштувань».",
            "Після підключення до Wi-Fi налаштуйте MQTT, перевірте web-панель та за потреби виконайте калібрування LED.",
        ],
        style_name="List Number",
    )

    add_section_lead(
        doc,
        "ОНОВЛЕННЯ ПРИСТРОЮ",
        "4. Прошивка через AlarmMini Installer",
        "Веб-інсталятор допомагає безпечно оновити плату, зчитати службову інформацію та, за потреби, відновити попередні налаштування після перепрошивки.",
    )
    add_text(
        doc,
        "Веб-інсталятор дозволяє підключити плату, прочитати службову інформацію, зберегти поточний JSON-конфіг, прошити пристрій і автоматично відновити налаштування після оновлення.",
        after=6,
    )
    add_three_col_table(
        doc,
        [
            ("1", "Підключення плати", "Натисніть «Підключити плату», оберіть серійний порт та дочекайтесь статусу «Порт підключено»."),
            ("2", "Службова інформація", "Натисніть «Оновити інформацію», щоб побачити FW, IP, mDNS, hostname, admin пароль, AP SSID та причину перезавантаження."),
            ("3", "Вибір режиму", "Для нової плати використовуйте «Прошити як новий пристрій». Для оновлення вже налаштованої плати використовуйте «Прошити з відновленням налаштувань»."),
            ("4", "Контроль процесу", "Слідкуйте за індикаторами Backup конфігу, Прошивка, Перепідключення, Відновлення Wi-Fi, Відновлення JSON та Перевірка."),
        ],
        header=("Крок", "Розділ", "Що робити"),
    )
    add_note(
        doc,
        "Порада",
        "Перед будь-яким оновленням натисніть «Зчитати конфігурацію» та збережіть резервну копію JSON. Це найпростіший спосіб швидко відновити робочі налаштування.",
    )

    add_heading(doc, "5. Налаштування в інсталяторі", 1)
    add_heading(doc, "5.1. Wi-Fi та MQTT", 2)
    add_bullets(
        doc,
        [
            "У вкладці «Мережа» можна окремо зберегти Wi-Fi та MQTT параметри без повної перепрошивки.",
            "Для Wi-Fi заповніть SSID і пароль та натисніть «Зберегти Wi-Fi».",
            "Для MQTT заповніть Host, Port, Topic, Username і Password та натисніть «Зберегти MQTT».",
        ],
    )
    add_heading(doc, "5.2. Config JSON", 2)
    add_bullets(
        doc,
        [
            "Редактор JSON призначений для розширеного налаштування або відновлення резервної копії.",
            "Кнопка «Перевірити сумісність конфігу» допомагає знайти помилки перед записом.",
            "Кнопки «Завантажити backup JSON», «Safe restore Wi-Fi+MQTT» та «Відновити backup JSON» корисні після невдалого редагування або оновлення.",
        ],
    )
    add_note(
        doc,
        "Обережно з JSON",
        "Якщо ви не впевнені, не змінюйте невідомі поля вручну. Найчастіше для звичайного користування достатньо налаштувати Wi-Fi, MQTT і калібрування через інтерфейс.",
    )

    add_section_lead(
        doc,
        "КЕРУВАННЯ ПРИСТРОЄМ",
        "6. Доступ до web-панелі пристрою",
        "Після отримання IP-адреси карту можна налаштовувати вже через вбудовану web-панель: кольори, нічний режим, MQTT, QR-коди та прив'язку LED.",
    )
    add_text(
        doc,
        "Після того як плата отримала IP-адресу, пристрій можна відкривати у браузері для подальшого налаштування карти, кольорів і логіки роботи.",
        after=6,
    )
    add_bullets(
        doc,
        [
            "Найпростіше відкрити панель через QR-код або кнопку «Відкрити web panel» в інсталяторі.",
            "Якщо mDNS працює у вашій мережі, адреса матиме вигляд `http://hostname.local`.",
            "Якщо mDNS недоступний, відкрийте пристрій за IP-адресою, яку показує інсталятор або роутер.",
            "Для входу використовується admin пароль, який генерується пристроєм і відображається в інсталяторі та на QR-наклейці.",
        ],
    )
    add_two_col_table(
        doc,
        [
            ("Огляд", "Поточний стан карти та коротка інформація про пристрій."),
            ("QR / Друк", "Генерація і друк QR-кодів для web-панелі та точки доступу."),
            ("Кольори", "Денні та нічні кольори для режимів «тривога» та «відбій»."),
            ("Нічний режим", "Час початку і завершення нічного профілю."),
            ("Бузер", "Увімкнення звуку та регулювання гучності. Доступність залежить від плати та збірки."),
            ("Регіони", "Вибір областей, для яких активні звукові сповіщення."),
            ("MQTT", "Адреса брокера, порт, topic, логін і пароль."),
            ("Система", "Admin пароль, mDNS, IP, службові параметри та NTP-сервери."),
            ("Калібрування", "Прив'язка номерів LED до областей України."),
        ],
    )

    add_section_lead(
        doc,
        "ТОЧНІСТЬ ВІДОБРАЖЕННЯ",
        "7. Калібрування LED",
        "Калібрування потрібне, якщо фізичне розташування світлодіодів у корпусі не збігається з тим, як області вже прив'язані в пам'яті пристрою.",
    )
    add_text(
        doc,
        "Калібрування потрібне, якщо фізичне розташування світлодіодів у вашій карті не збігається з поточною мапою областей у пам'яті пристрою.",
        after=6,
    )
    add_bullets(
        doc,
        [
            "Відкрийте розділ «Калібрування» у web-панелі.",
            "Натисніть «Почати» або «Відкрити калібрування».",
            "Оберіть LED крок за кроком або скористайтеся швидким переходом по LED.",
            "Для кожного світлодіода виберіть область на карті або зі списку областей.",
            "На фінальному екрані перевірте підсумок і натисніть «Зберегти мапу».",
        ],
    )
    add_note(
        doc,
        "Практична порада",
        "Починайте калібрування з першого фізичного LED на стрічці й рухайтесь у напрямку її підключення. Так легше не збитися при зіставленні областей.",
    )

    add_heading(doc, "8. QR-коди та наклейки", 1)
    add_bullets(
        doc,
        [
            "Admin QR веде на web-панель пристрою і підходить для власника карти.",
            "AP QR призначений для швидкого входу в режим налаштування Wi-Fi, коли пристрій працює як точка доступу.",
            "У інсталяторі можна завантажити кожен QR-код у PNG або відправити їх на друк.",
        ],
    )

    add_heading(doc, "9. Що означає поведінка пристрою", 1)
    add_two_col_table(
        doc,
        [
            ("Пристрій створив Wi-Fi AlarmMap-Setup", "Ще не збережено Wi-Fi або пристрій не зміг підключитися до мережі."),
            ("Сторінка 192.168.4.1 відкривається", "Плата працює в режимі налаштування точки доступу."),
            ("В інсталяторі видно IP, hostname та admin пароль", "Базовий зв'язок із пристроєм працює, можна переходити до web-панелі."),
            ("Після прошивки з'являється статус перевірки", "Інсталятор відновлює конфіг і звіряє результат."),
        ],
    )

    add_heading(doc, "10. Поширені проблеми", 1)
    add_three_col_table(
        doc,
        [
            ("Плата не підключається в інсталяторі", "Кабель лише для зарядки, несумісний браузер, не той порт", "Спробуйте інший USB-кабель, використайте Chrome/Edge і перевиберіть серійний порт."),
            ("Не підключається до Wi-Fi", "Неправильний пароль або мережа 5 ГГц", "Перевірте пароль і використовуйте мережу 2.4 ГГц."),
            ("Web-панель не відкривається", "Пристрій і комп'ютер у різних мережах або змінилась IP-адреса", "Перевірте IP у роутері чи в інсталяторі та відкривайте пристрій у тій самій локальній мережі."),
            ("LED не реагують на тривоги", "Некоректні MQTT дані або не виконано калібрування", "Перевірте Host, Port, Topic та прив'язку LED до областей."),
            ("Нічний режим працює неправильно", "Не синхронізувався час або задано помилковий діапазон", "Перевірте NTP-сервери, доступ до інтернету і час початку/завершення нічного режиму."),
        ],
        header=("Проблема", "Ймовірна причина", "Що зробити"),
    )

    add_heading(doc, "11. Рекомендації з безпечної експлуатації", 1)
    add_bullets(
        doc,
        [
            "Перед оновленням прошивки завжди робіть резервну копію конфігурації.",
            "Не від'єднуйте USB-кабель під час активної прошивки або відновлення JSON.",
            "Зберігайте admin пароль і QR-наклейки в доступному місці, якщо карта передається іншому користувачу.",
            "Якщо змінюєте MQTT або калібрування вручну, після збереження перевірте карту на реальному оновленні стану.",
        ],
    )
    add_note(
        doc,
        "Зв'язок з автором",
        "Якщо потрібна допомога, автору можна написати в Telegram: " + TELEGRAM_URL,
    )

    add_text(
        doc,
        "Актуальний інсталятор: https://alarmmini.vercel.app\n"
        "Репозиторій проєкту: https://github.com/WebDev-Den/AlarmMini",
        size=10,
        color=GRAY,
        before=16,
        after=0,
        line=1.2,
    )
    return doc


if __name__ == "__main__":
    document = build_doc()
    document.save(DOCX_PATH)
    print(DOCX_PATH)
